from datetime import datetime
from pathlib import Path
import tempfile
import unittest
import zipfile

from docx import Document

from tendertrace.acceptance import REQUIRED_DELIVERY_DOCS, run_acceptance
from tendertrace.config import Settings
from tendertrace.db import connection, init_db, json_dumps
from tendertrace.runlog import finish_run, register_outbox_message, start_run


class AcceptanceTests(unittest.TestCase):
    def test_acceptance_passes_for_complete_workspace_fixture(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_required_files(root)
            settings = Settings.load(root)
            init_db(settings)
            docx_path = _write_report(settings.outputs_dir / "示例问题_202607061300.docx")
            settings.outbox_dir.mkdir(parents=True, exist_ok=True)
            outbox_path = settings.outbox_dir / docx_path.name
            outbox_path.write_bytes(docx_path.read_bytes())
            _insert_finished_run(settings, outbox_path)
            with connection(settings) as conn:
                conn.execute(
                    """
                    INSERT INTO model_audits(
                        id, run_id, mode, provider, model, status, prompt_sha256, response_sha256
                    )
                    VALUES ('audit-1', 'run-1', 'cloud', 'openai', 'test-model', 'ok', ?, ?)
                    """,
                    ("a" * 64, "b" * 64),
                )

            report = run_acceptance(settings)

        failed = [check.to_dict() for check in report.checks if check.status == "fail"]
        self.assertEqual(report.status, "pass")
        self.assertEqual(failed, [])

    def test_acceptance_fails_when_env_example_contains_plaintext_key(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_required_files(root)
            (root / ".env.example").write_text(
                "OPENAI_API_KEY=sk-proj-" + "A" * 32,
                encoding="utf-8",
            )
            settings = Settings.load(root)
            init_db(settings)

            report = run_acceptance(settings, strict_runtime=False)

        failed_names = {check.name for check in report.checks if check.status == "fail"}
        self.assertEqual(report.status, "fail")
        self.assertIn("env_example", failed_names)
        self.assertIn("secret_scan", failed_names)

    def test_acceptance_accepts_multi_source_attempt_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_required_files(root)
            settings = Settings.load(root)
            init_db(settings)
            docx_path = _write_report(settings.outputs_dir / "示例问题_202607061300.docx")
            settings.outbox_dir.mkdir(parents=True, exist_ok=True)
            outbox_path = settings.outbox_dir / docx_path.name
            outbox_path.write_bytes(docx_path.read_bytes())
            _insert_finished_run(
                settings,
                outbox_path,
                source_sites=["ggzy"],
                source_stats=[
                    {"source": "local_fts5", "status": "finished", "count": 1},
                    {"source": "ccgp", "status": "finished", "count": 0},
                    {"source": "ggzy", "status": "finished", "count": 1},
                ],
            )

            report = run_acceptance(settings)

        checks = {check.name: check for check in report.checks}
        self.assertEqual(checks["multi_source_run"].status, "pass")
        self.assertIn("attempted ccgp, ggzy", checks["multi_source_run"].detail)

    def test_acceptance_reports_invalid_submission_zip_without_crashing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _write_required_files(root)
            (root / "dist" / "TenderTrace_submission_202607062130.zip").write_bytes(
                b"not-a-zip" * 2000
            )
            settings = Settings.load(root)
            init_db(settings)

            report = run_acceptance(settings, strict_runtime=False)

        checks = {check.name: check for check in report.checks}
        self.assertEqual(report.status, "fail")
        self.assertEqual(checks["submission_package"].status, "fail")
        self.assertIn("could not be scanned", checks["submission_package"].detail)


def _write_required_files(root: Path) -> None:
    files = [
        ".env.example",
        *REQUIRED_DELIVERY_DOCS,
        "docs/teaching/13_交付收口与完成度审计.docx",
        "docs/demo/demo演示视频.mp4",
        "dist/TenderTrace_submission_202607062130.zip",
    ]
    for relative in files:
        path = root / relative
        path.parent.mkdir(parents=True, exist_ok=True)
        if path.suffix == ".docx":
            _write_report(path)
        elif path.suffix == ".mp4":
            path.write_bytes(b"0" * 20000)
        elif path.suffix == ".zip":
            _write_submission_zip(path)
        elif path.name == ".env.example":
            path.write_text("OPENAI_API_KEY=\n", encoding="utf-8")
        else:
            path.write_text(f"# {relative}\n\n" + "content\n" * 80, encoding="utf-8")


def _write_report(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("标题：示例公告")
    document.add_paragraph("发布时间：2026-07-06")
    document.add_paragraph("来源链接：https://example.com")
    document.add_paragraph("核心内容：示例采购项目")
    document.save(path)
    return path


def _write_submission_zip(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("README.md", "content" * 3000)
        archive.writestr("SUBMISSION_MANIFEST.json", '{"file_count": 1}')


def _insert_finished_run(
    settings: Settings,
    outbox_path: Path,
    *,
    source_sites: list[str] | None = None,
    source_stats: list[dict[str, object]] | None = None,
) -> None:
    started = datetime(2026, 7, 6, 13, 0).isoformat()
    start_run(
        settings,
        run_id="run-1",
        original_query="示例问题",
        mode="full",
        window_start=started,
        window_end=started,
    )
    finish_run(
        settings,
        run_id="run-1",
        status="finished",
        output_docx_path=outbox_path,
        stats={
            "notice_count": 1,
            "trace_events": 7,
            "evidence_checked": 1,
            "source_sites": source_sites or ["ccgp", "ggzy"],
            "source_stats": source_stats or [],
        },
    )
    register_outbox_message(settings, run_id="run-1", docx_path=outbox_path)
    with connection(settings) as conn:
        conn.execute(
            """
            INSERT INTO trace_events(run_id, seq, event_type, node, payload_json)
            VALUES ('run-1', 1, 'tool_called', 'intent', ?)
            """,
            (json_dumps({"tool": "intent.rule_parser"}),),
        )


if __name__ == "__main__":
    unittest.main()
