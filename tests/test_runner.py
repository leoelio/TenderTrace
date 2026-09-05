from datetime import datetime
from io import BytesIO
from pathlib import Path
import tempfile
import unittest
from unittest.mock import patch
from zoneinfo import ZoneInfo

from docx import Document

from tendertrace.adapters.ccgp import Attachment, Notice
from tendertrace.adapters.multi import MultiSourceAdapter
from tendertrace.config import Settings
from tendertrace.db import connection
from tendertrace.delivery.feishu_report import FeishuReportDelivery
from tendertrace.llm.gateway import ModelCallResult
from tendertrace.runlog import get_run, list_outbox_messages
from tendertrace.runner import _can_use_local_only, run_once
from tendertrace.runtime.checkpoint import SqliteCheckpointer
from tendertrace.runtime.trace import SqliteTraceStore


class FakeAdapter:
    def __init__(self) -> None:
        self.calls: list[dict[str, object]] = []

    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        self.calls.append({"bidql": bidql, "max_pages": max_pages, "max_results": max_results})
        return [
            Notice(
                id="notice-1",
                source_site="ccgp",
                title="上海某单位服务器采购公开招标公告",
                publish_time="2026-07-06 09:30",
                region="上海",
                purchaser="上海某单位",
                source_url="https://www.ccgp.gov.cn/example.htm",
                core_content="项目概况：上海某单位服务器采购项目。预算金额：120万元。",
            )
        ]


class ChargingFacilityAdapter:
    def __init__(self) -> None:
        self.calls = 0

    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        self.calls += 1
        return [
            Notice(
                id="charging-1",
                source_site="ccgp",
                title="上海充电设施建设项目公开招标公告",
                publish_time="2026-07-01 09:00",
                region="上海",
                purchaser="上海某单位",
                source_url="https://www.ccgp.gov.cn/charging-1.html",
                content_text="采购内容为新能源汽车充电基础设施。",
                core_content="采购内容为新能源汽车充电基础设施。",
            )
        ]


class FailingAdapter:
    def __init__(self) -> None:
        self.calls = 0

    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        self.calls += 1
        raise AssertionError("local retrieval should satisfy this run")


class EbrdOnlyAdapter(FailingAdapter):
    name = "ebrd"

    def supports(self, bidql: dict[str, object]) -> bool:
        return bidql.get("region", {}).get("scope") == "ebrd"


class SupplementAdapter:
    def __init__(self) -> None:
        self.calls = 0

    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        self.calls += 1
        return [
            Notice(
                id="ggzy-charging-2",
                source_site="ggzy",
                title="上海充电桩运维服务招标公告",
                publish_time="2026-07-02 10:00",
                region="上海",
                purchaser="上海公共资源交易中心",
                source_url="https://www.ggzy.gov.cn/charging-2.html",
                content_text="采购内容为充电桩运维服务。",
                core_content="采购内容为充电桩运维服务。",
            )
        ]


class DuplicateAdapter:
    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        return [
            Notice(
                id="ccgp-1",
                source_site="ccgp",
                title="上海某医院医疗设备采购项目公开招标公告",
                publish_time="2026-07-06 09:30",
                region="上海",
                purchaser="上海某医院",
                source_url="https://www.ccgp.gov.cn/detail.html",
                content_text="项目编号：SH-2026-001。项目概况：医疗设备采购。打印本页",
                core_content="项目概况：医疗设备采购。",
            ),
            Notice(
                id="ggzy-1",
                source_site="ggzy",
                title="上海某医院医疗设备采购项目招标公告",
                publish_time="2026-07-06",
                region="上海市",
                purchaser="上海政府采购网",
                source_url="https://www.ggzy.gov.cn/detail.html",
                content_text="项目编号：SH-2026-001。项目概况：医疗设备采购。",
                core_content="项目概况：医疗设备采购。",
            ),
        ]


class AttachmentAdapter:
    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        return [
            Notice(
                id="attachment-1",
                source_site="ccgp",
                title="上海服务器采购公开招标公告",
                publish_time="2026-07-06 09:30",
                region="上海",
                purchaser="上海某单位",
                source_url="https://www.ccgp.gov.cn/attachment-detail.html",
                content_text="项目编号：SH-2026-ATT。项目概况：上海服务器采购。",
                core_content="项目概况：上海服务器采购。",
                attachments=[
                    Attachment(name="采购需求", url="https://example.com/spec.docx"),
                ],
            )
        ]


class ArtifactAdapter:
    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        return [
            Notice(
                id="artifact-1",
                source_site="ccgp",
                title="Shanghai server tender notice",
                publish_time="2026-07-06 09:30",
                region="Shanghai",
                purchaser="Shanghai buyer",
                source_url="https://www.ccgp.gov.cn/artifact-detail.html",
                content_text="Project No: SH-2026-ART. Project scope: server purchase.",
                core_content="Project scope: server purchase.",
                fields={
                    "page_artifact": {
                        "source_site": "ccgp",
                        "source_url": "https://www.ccgp.gov.cn/artifact-detail.html",
                        "final_url": "https://www.ccgp.gov.cn/artifact-detail.html",
                        "status_code": 200,
                        "fetcher": "httpx",
                        "content_sha256": "abc123",
                        "content_length": 128,
                        "text_excerpt": "detail text",
                        "blocked": False,
                        "error": "",
                        "fetched_at": "2026-07-06T02:00:00+00:00",
                        "elapsed_ms": 12,
                    }
                },
            )
        ]


class RelaxedCityAdapter:
    def collect(
        self,
        bidql: dict[str, object],
        *,
        max_pages: int = 1,
        max_results: int = 10,
    ) -> list[Notice]:
        self.last_source_stats = [
            {
                "source": "ggzy",
                "status": "finished",
                "count": 1,
                "error": None,
                "relaxed_city": True,
            }
        ]
        return [
            Notice(
                id="zj-1",
                source_site="ggzy",
                title="浙江省服务器采购项目招标公告",
                publish_time="2026-07-06",
                region="浙江省",
                purchaser="浙江省公共资源交易中心",
                source_url="https://www.ggzy.gov.cn/zj-1.html",
                content_text="项目概况：浙江省服务器采购。",
                core_content="项目概况：浙江省服务器采购。",
            )
        ]


class FakeModelGateway:
    def __init__(self) -> None:
        self.calls: list[dict[str, str]] = []

    def generate_json(self, *, system: str, user: str) -> ModelCallResult:
        self.calls.append({"system": system, "user": user})
        return ModelCallResult(
            mode="local",
            provider="ollama",
            model="qwen3:8b",
            status="ok",
            text='{"expanded_keywords":["GPU服务器"],"negative_keywords":["废标"]}',
            parsed={
                "expanded_keywords": ["GPU服务器"],
                "negative_keywords": ["废标"],
                "intent_summary": "server tender search",
                "confidence": 0.89,
            },
            latency_ms=12,
        )


class RunnerTests(unittest.TestCase):
    def test_run_once_writes_report_outbox_trace_and_runlog(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))
            adapter = FakeAdapter()

            result = run_once(
                settings=settings,
                query="最近1个月的上海区域内的服务器招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                max_pages=2,
                max_results=5,
                adapter=adapter,
            )

            run = get_run(settings, result.run_id)
            messages = list_outbox_messages(settings)
            checkpoints = SqliteCheckpointer(settings).list(result.run_id)
            events = SqliteTraceStore(settings).list_events(result.run_id)
            docx_exists = Path(result.docx_path).exists()
            outbox_exists = Path(result.outbox_path).exists()

        self.assertEqual(result.status, "finished")
        self.assertEqual(result.notice_count, 1)
        self.assertTrue(docx_exists)
        self.assertTrue(outbox_exists)
        self.assertEqual(run["stats"]["notice_count"], 1)
        self.assertEqual(messages[0].run_id, result.run_id)
        self.assertEqual(
            [checkpoint.node for checkpoint in checkpoints],
            ["intent", "collect", "evidence", "report"],
        )
        self.assertGreaterEqual(len(events), 4)
        self.assertEqual(run["stats"]["evidence_checked"], 1)
        self.assertEqual(run["stats"]["feishu_bitable_delivery"]["status"], "skipped")
        self.assertEqual(adapter.calls[0]["max_pages"], 2)
        self.assertEqual(adapter.calls[0]["max_results"], 5)
        self.assertTrue(
            any(event.payload.get("tool") == "delivery.feishu_bitable" for event in events)
        )

    def test_run_once_sends_evidence_summary_with_feishu_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))
            delivery = FeishuReportDelivery(
                status="sent",
                file_name="report.docx",
                attempt_id="attempt-1",
                digest_status="sent",
            )
            with patch("tendertrace.runner.deliver_report_to_feishu", return_value=delivery) as sender:
                run_once(
                    settings=settings,
                    query="最近1个月上海服务器招标信息都有哪些",
                    now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                    adapter=FakeAdapter(),
                    delivery_channels=("web", "outbox", "feishu"),
                )

        summary = sender.call_args.kwargs["report_summary"]
        self.assertEqual(summary["query"], "最近1个月上海服务器招标信息都有哪些")
        self.assertEqual(summary["notice_count"], 1)
        self.assertIn("ccgp", summary["source_sites"])
        self.assertEqual(summary["highlights"][0]["title"], "上海某单位服务器采购公开招标公告")

    def test_run_once_applies_model_enhancement_and_records_audit(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / ".env.local").write_text(
                "TENDERTRACE_MODEL_ENHANCEMENT_ENABLED=true\n",
                encoding="utf-8",
            )
            settings = Settings.load(root)
            adapter = FakeAdapter()
            model_gateway = FakeModelGateway()

            result = run_once(
                settings=settings,
                query="最近1个月上海服务器招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                adapter=adapter,
                model_gateway=model_gateway,
            )

            events = SqliteTraceStore(settings).list_events(result.run_id)
            with connection(settings) as conn:
                audit_rows = conn.execute("SELECT * FROM model_audits").fetchall()

        bidql = adapter.calls[0]["bidql"]
        expanded_terms = [
            item["term"] for item in bidql["topic"]["expanded"] if isinstance(item, dict)
        ]
        self.assertEqual(model_gateway.calls[0]["system"][:11], "You enhance")
        self.assertIn("GPU服务器", expanded_terms)
        self.assertIn("废标", bidql["topic"]["negative"])
        self.assertGreaterEqual(len(audit_rows), 1)
        self.assertEqual(audit_rows[0]["status"], "ok")
        self.assertTrue(any(event.payload.get("tool") == "llm.intent_enhancer" for event in events))

    def test_run_once_deduplicates_and_persists_clusters(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))

            result = run_once(
                settings=settings,
                query="最近1个月上海医疗设备招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                adapter=DuplicateAdapter(),
            )

            run = get_run(settings, result.run_id)
            with connection(settings) as conn:
                cluster_rows = conn.execute("SELECT * FROM clusters").fetchall()
                notice_rows = conn.execute("SELECT * FROM notices").fetchall()
                evidence_rows = conn.execute("SELECT * FROM evidence_items").fetchall()

        self.assertEqual(result.notice_count, 1)
        self.assertEqual(run["stats"]["raw_count"], 2)
        self.assertEqual(run["stats"]["duplicates_removed"], 1)
        self.assertEqual(run["stats"]["cluster_count"], 1)
        self.assertEqual(run["stats"]["source_sites"], ["ccgp", "ggzy"])
        self.assertEqual(len(cluster_rows), 1)
        self.assertEqual(cluster_rows[0]["cluster_key"], "project:sh-2026-001")
        self.assertEqual(len(notice_rows), 1)
        self.assertEqual(len(evidence_rows), 1)
        self.assertEqual(evidence_rows[0]["cluster_key"], "project:sh-2026-001")

    def test_run_once_reuses_local_fts_before_collecting_again(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))
            seed_adapter = ChargingFacilityAdapter()
            first = run_once(
                settings=settings,
                query="最近1个月上海充电设施招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                max_results=1,
                adapter=seed_adapter,
            )
            failing_adapter = FailingAdapter()
            second = run_once(
                settings=settings,
                query="最近1个月上海充电桩招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                max_results=1,
                adapter=failing_adapter,
            )

        self.assertEqual(first.notice_count, 1)
        self.assertEqual(seed_adapter.calls, 1)
        self.assertEqual(second.notice_count, 1)
        self.assertEqual(failing_adapter.calls, 0)
        self.assertEqual(second.stats["local_retrieved"], 1)
        self.assertEqual(second.stats["source_collected"], 0)
        self.assertEqual(second.stats["retrieval_engine"], "fts5")

    def test_run_once_keeps_live_collection_when_local_cache_has_one_source(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))
            seed_adapter = ChargingFacilityAdapter()
            run_once(
                settings=settings,
                query="最近1个月上海充电设施招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                max_results=1,
                adapter=seed_adapter,
            )
            supplement = SupplementAdapter()

            second = run_once(
                settings=settings,
                query="最近1个月上海充电桩招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                max_results=1,
                adapter=MultiSourceAdapter([supplement]),
            )

        self.assertEqual(supplement.calls, 1)
        self.assertEqual(second.stats["local_retrieved"], 1)
        self.assertEqual(second.stats["source_collected"], 1)
        self.assertEqual(second.stats["source_sites"], ["ccgp", "ggzy"])

    def test_scoped_single_source_can_reuse_a_full_local_result(self) -> None:
        notice = Notice(
            id="ebrd-water-1",
            source_site="ebrd",
            title="Water infrastructure procurement",
            publish_time="2026-08-13",
            region="Jordan",
            purchaser="Water Authority",
            source_url="https://ecepp.ebrd.com/delta/viewNotice.html?displayNoticeId=1",
        )
        adapter = MultiSourceAdapter([EbrdOnlyAdapter()])

        self.assertTrue(
            _can_use_local_only(
                [notice],
                1,
                adapter,
                {"region": {"scope": "ebrd"}},
            )
        )

    def test_run_once_downloads_extracts_and_persists_attachment_snapshots(self) -> None:
        def fake_download(url: str, max_bytes: int) -> bytes:
            self.assertEqual(url, "https://example.com/spec.docx")
            return _docx_bytes("附件正文：服务器配置要求和交付周期。")

        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))

            result = run_once(
                settings=settings,
                query="最近1个月上海服务器招标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                adapter=AttachmentAdapter(),
                attachment_downloader=fake_download,
            )

            run = get_run(settings, result.run_id)
            report = Document(result.docx_path)
            report_text = "\n".join(paragraph.text for paragraph in report.paragraphs)
            with connection(settings) as conn:
                attachment_rows = conn.execute("SELECT * FROM attachment_snapshots").fetchall()

        self.assertEqual(result.status, "finished")
        self.assertEqual(run["stats"]["attachments_extracted"], 1)
        self.assertEqual(len(attachment_rows), 1)
        self.assertEqual(attachment_rows[0]["status"], "extracted")
        self.assertIn("附件解析：已下载 1/1，已抽取正文 1，失败 0，跳过 0", report_text)
        self.assertIn("附件正文：服务器配置要求和交付周期。", report_text)

    def test_run_once_persists_page_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))

            result = run_once(
                settings=settings,
                query="server tender",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                adapter=ArtifactAdapter(),
            )

            with connection(settings) as conn:
                rows = conn.execute("SELECT * FROM page_artifacts").fetchall()

        self.assertEqual(result.status, "finished")
        self.assertGreaterEqual(result.stats["structured_field_hits"], 4)
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["source_site"], "ccgp")
        self.assertEqual(rows[0]["status_code"], 200)
        self.assertEqual(rows[0]["content_sha256"], "abc123")

    def test_run_once_records_city_scope_relaxation_in_stats_and_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))

            result = run_once(
                settings=settings,
                query="最近36个月杭州市的空调或者服务器投标信息都有哪些",
                now=datetime(2026, 7, 6, 10, 0, tzinfo=ZoneInfo("Asia/Shanghai")),
                adapter=RelaxedCityAdapter(),
            )

            run = get_run(settings, result.run_id)
            report = Document(result.docx_path)
            report_text = "\n".join(paragraph.text for paragraph in report.paragraphs)

        self.assertEqual(run["stats"]["region_scope"]["status"], "relaxed_city")
        self.assertEqual(run["stats"]["region_scope"]["requested_city"], "杭州")
        self.assertIn("杭州城市级检索未命中样本，已扩大到浙江省内检索", report_text)
        self.assertIn("地区：杭州（城市级无结果，已扩大至浙江省内）", report_text)


def _docx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


if __name__ == "__main__":
    unittest.main()
