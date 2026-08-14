from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from tendertrace.adapters.ccgp import Attachment, Notice
from tendertrace.report.docx_writer import write_report
from tendertrace.report.naming import safe_report_filename


class ReportWriterTests(unittest.TestCase):
    def test_safe_report_filename_keeps_query_and_timestamp(self) -> None:
        filename = safe_report_filename(
            "最近1个月的上海区域内的服务器招标信息都有哪些",
            datetime(2026, 7, 6, 14, 24),
        )

        self.assertEqual(
            filename, "最近1个月的上海区域内的服务器招标信息都有哪些_202607061424.docx"
        )

    def test_write_report_contains_required_notice_fields(self) -> None:
        bidql = {
            "topic": {"core": ["服务器"]},
            "region": {"province": "上海"},
            "time": {"resolved_window": {"from": "2026-06-06", "to": "2026-07-06"}},
        }
        notice = Notice(
            id="n1",
            source_site="ccgp",
            title="上海某单位服务器采购公开招标公告",
            publish_time="2026-07-06 09:30",
            region="上海",
            purchaser="上海某单位",
            source_url="https://www.ccgp.gov.cn/cggg/dfgg/gkzb/202607/t20260706_123.htm",
            core_content="项目概况：上海某单位服务器采购公开招标项目。预算金额：120万元。",
            attachments=[
                Attachment(
                    name="附件：采购需求",
                    url="https://www.ccgp.gov.cn/cggg/dfgg/files/spec.docx",
                )
            ],
            fields={
                "evidence": {
                    "status": "passed",
                    "quality_score": 1.0,
                    "snapshot_sha256": "a" * 64,
                    "excerpt": "项目概况：上海某单位服务器采购公开招标项目。预算金额：120万元。",
                    "attachments": [
                        {
                            "name": "附件：采购需求",
                            "url": "https://www.ccgp.gov.cn/cggg/dfgg/files/spec.docx",
                            "type": "docx",
                            "status": "extracted",
                            "path": "snapshots/attachments/spec.docx",
                            "sha256": "b" * 64,
                            "bytes": 1024,
                            "text_excerpt": "附件正文：服务器配置要求和交付周期。",
                            "text_length": 18,
                            "error": "",
                        }
                    ],
                }
            },
        )

        with TemporaryDirectory() as tmp:
            path = write_report(
                query="最近1个月的上海区域内的服务器招标信息都有哪些",
                bidql=bidql,
                notices=[notice],
                output_dir=Path(tmp),
                generated_at=datetime(2026, 7, 6, 14, 24),
                run_stats={
                    "collected": 3,
                    "local_retrieved": 1,
                    "source_collected": 2,
                    "deduped": 1,
                    "source_stats": [
                        {
                            "source": "ccgp",
                            "status": "finished",
                            "count": 1,
                            "fetch_stats": {
                                "requests": 2,
                                "successes": 2,
                                "blocked": 0,
                                "retries": 1,
                            },
                        },
                        {
                            "source": "ggzy",
                            "status": "finished",
                            "count": 0,
                            "relaxed_city": True,
                            "fetch_stats": {"requests": 1, "successes": 1},
                        },
                    ],
                },
            )

            doc = Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(
                cell.text for table in doc.tables for row in table.rows for cell in row.cells
            )

        self.assertTrue(path.name.endswith("_202607061424.docx"))
        self.assertIn("TenderTrace 招投标信息汇总报告", text)
        self.assertIn(
            "运行漏斗：候选 3 条，本地复用 1 条，外部来源新增 2 条，清洗去重后 1 条。", text
        )
        self.assertIn("多源覆盖：本轮尝试 2 个来源，1 个来源命中，0 个来源异常。", text)
        self.assertIn("来源覆盖与抓取健康", text)
        self.assertIn("机会优先级", text)
        self.assertIn("市场研判", text)
        self.assertIn("市场价格位置", text)
        self.assertIn("竞争情报", text)
        self.assertIn("需求覆盖", text)
        self.assertIn("需求待核对", text)
        self.assertIn("需求研判边界", text)
        self.assertIn("ccgp", table_text)
        self.assertIn("城市无结果，已放宽到省级", table_text)
        self.assertIn("标题：上海某单位服务器采购公开招标公告", text)
        self.assertIn("发布时间：2026-07-06 09:30", text)
        self.assertIn(
            "来源链接：https://www.ccgp.gov.cn/cggg/dfgg/gkzb/202607/t20260706_123.htm", text
        )
        self.assertIn(
            "核心内容：项目概况：上海某单位服务器采购公开招标项目。预算金额：120万元。", text
        )
        self.assertIn("事实校验：passed（score: 1.0）", text)
        self.assertIn(
            "证据摘录：项目概况：上海某单位服务器采购公开招标项目。预算金额：120万元。", text
        )
        self.assertIn("附件解析：已下载 1/1，已抽取正文 1，失败 0，跳过 0", text)
        self.assertIn("附件正文：服务器配置要求和交付周期。", text)
        self.assertIn("机会研判：", text)
        self.assertIn("质量分项：", text)
        self.assertIn("项目目标：", text)
        self.assertIn("附件：采购需求 - https://www.ccgp.gov.cn/cggg/dfgg/files/spec.docx", text)


if __name__ == "__main__":
    unittest.main()
