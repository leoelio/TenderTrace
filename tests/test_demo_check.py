from pathlib import Path
import json
import tempfile
import unittest

from docx import Document

from tendertrace.config import Settings
from tendertrace.db import connection, init_db, json_dumps
from tendertrace.demo_check import run_demo_check, write_demo_evidence
from tendertrace.runlog import finish_run, register_outbox_message, start_run
from tendertrace.submission import create_submission_package


class DemoCheckTests(unittest.TestCase):
    def test_demo_check_passes_with_run_trace_outbox_and_incremental_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            settings = Settings.load(root)
            init_db(settings)
            first_docx = _write_report(settings.outputs_dir / "问题一_202607061300.docx")
            second_docx = _write_report(settings.outputs_dir / "问题二_202607061301.docx")
            settings.outbox_dir.mkdir(parents=True, exist_ok=True)
            first_outbox = settings.outbox_dir / first_docx.name
            second_outbox = settings.outbox_dir / second_docx.name
            first_outbox.write_bytes(first_docx.read_bytes())
            second_outbox.write_bytes(second_docx.read_bytes())
            _insert_run(settings, "run-1", "最近7天全国服务器招标信息", first_outbox)
            _insert_run(settings, "run-2", "2026年3月上海充电桩招标信息", second_outbox)
            _insert_demo_trace(settings, "run-2")
            with connection(settings) as conn:
                conn.execute(
                    "UPDATE runs SET finished_at = '2026-07-06 13:00:00' WHERE id = 'run-1'"
                )
                conn.execute(
                    "UPDATE runs SET finished_at = '2026-07-06 13:01:00' WHERE id = 'run-2'"
                )
            _insert_subscription_and_sent_history(settings, "run-2", second_outbox)
            _write_text(root / ".github" / "workflows" / "ci.yml", "name: CI\n")
            create_submission_package(root)

            report = run_demo_check(settings)
            out = write_demo_evidence(report, root / "docs" / "demo" / "evidence.json")
            written_payload = json.loads(out.read_text(encoding="utf-8"))
            out_exists = out.exists()

        checks = {check.name: check for check in report.checks}
        self.assertEqual(report.status, "pass")
        self.assertEqual(checks["finished_runs"].status, "pass")
        self.assertEqual(checks["word_outbox"].status, "pass")
        self.assertEqual(checks["trace_flow"].status, "pass")
        self.assertEqual(checks["subscription_incremental"].status, "pass")
        self.assertEqual(checks["submission_package"].status, "pass")
        self.assertEqual(checks["ci_config"].status, "pass")
        self.assertEqual(checks["api_security"].status, "pass")
        self.assertEqual(checks["sources"].status, "warn")
        self.assertEqual(checks["demo_video_file"].status, "warn")
        self.assertTrue(out_exists)
        self.assertEqual(written_payload["status"], "pass")
        self.assertNotIn("fixture-token", json.dumps(written_payload, ensure_ascii=False))
        latest_stats = written_payload["evidence"]["latest_finished_run"]["stats"]
        self.assertEqual(latest_stats["feishu_bitable_delivery"]["app_token"], "[redacted]")

    def test_demo_check_fails_when_core_demo_evidence_is_missing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            settings = Settings.load(Path(tmp))
            init_db(settings)

            report = run_demo_check(settings)

        failed = {check.name for check in report.checks if check.status == "fail"}
        self.assertEqual(report.status, "fail")
        self.assertIn("finished_runs", failed)
        self.assertIn("word_outbox", failed)
        self.assertIn("trace_flow", failed)
        self.assertIn("subscription_incremental", failed)

    def test_demo_check_resolves_migrated_output_by_filename(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            settings = Settings.load(root)
            init_db(settings)
            report_path = _write_report(settings.outputs_dir / "迁移报告_202607061301.docx")
            settings.outbox_dir.mkdir(parents=True, exist_ok=True)
            outbox_path = settings.outbox_dir / report_path.name
            outbox_path.write_bytes(report_path.read_bytes())
            stale_path = Path("D:/old-workspace/outputs") / report_path.name
            _insert_run(settings, "run-migrated", "迁移后的查询", stale_path)

            report = run_demo_check(settings)

        checks = {check.name: check for check in report.checks}
        self.assertNotEqual(checks["word_outbox"].status, "fail")

    def test_demo_check_uses_latest_finished_run_with_existing_output(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            settings = Settings.load(root)
            init_db(settings)
            report_path = _write_report(settings.outputs_dir / "existing_202607061301.docx")
            settings.outbox_dir.mkdir(parents=True, exist_ok=True)
            (settings.outbox_dir / report_path.name).write_bytes(report_path.read_bytes())
            _insert_run(settings, "run-existing", "existing report", report_path)
            _insert_run(settings, "run-missing", "removed report", root / "removed.docx")
            with connection(settings) as conn:
                conn.execute(
                    "UPDATE runs SET finished_at = '2026-07-06 13:01:00' "
                    "WHERE id = 'run-existing'"
                )
                conn.execute(
                    "UPDATE runs SET finished_at = '2026-07-06 13:02:00' "
                    "WHERE id = 'run-missing'"
                )

            report = run_demo_check(settings)

        self.assertEqual(report.evidence["latest_finished_run"]["id"], "run-existing")
        checks = {check.name: check for check in report.checks}
        self.assertNotEqual(checks["word_outbox"].status, "fail")


def _write_report(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("标题：示例公告")
    doc.add_paragraph("发布时间：2026-07-06")
    doc.add_paragraph("来源链接：https://example.com")
    doc.add_paragraph("核心内容：示例采购项目")
    doc.save(path)
    return path


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def _insert_run(settings: Settings, run_id: str, query: str, docx_path: Path) -> None:
    start_run(
        settings,
        run_id=run_id,
        original_query=query,
        mode="full",
        window_start="2026-07-01",
        window_end="2026-07-06",
    )
    register_outbox_message(settings, run_id=run_id, docx_path=docx_path)
    finish_run(
        settings,
        run_id=run_id,
        status="finished",
        output_docx_path=docx_path,
        stats={
            "notice_count": 1,
            "trace_events": 7,
            "evidence_checked": 1,
            "source_sites": ["ccgp", "ggzy"],
            "feishu_bitable_delivery": {
                "status": "sent",
                "app_token": "fixture-token",
                "table_id": "tbl_test",
            },
        },
    )


def _insert_demo_trace(settings: Settings, run_id: str) -> None:
    tools = [
        "intent.rule_parser",
        "llm.intent_enhancer",
        "adapter.multi.collect",
        "pipeline.clean_dedup",
        "pipeline.attachment_extract",
        "pipeline.evidence_validate",
        "report.docx_writer",
    ]
    with connection(settings) as conn:
        for seq, tool in enumerate(tools, start=1):
            conn.execute(
                """
                INSERT INTO trace_events(run_id, seq, event_type, node, payload_json)
                VALUES (?, ?, 'tool_called', 'demo', ?)
                """,
                (run_id, seq, json_dumps({"tool": tool})),
            )


def _insert_subscription_and_sent_history(
    settings: Settings,
    run_id: str,
    docx_path: Path,
) -> None:
    with connection(settings) as conn:
        conn.execute(
            """
            INSERT INTO subscriptions(id, original_query, bidql_json, schedule_kind, cron, timezone)
            VALUES ('sub-1', '每天9点发送设备招标信息', ?, 'recurring', '0 9 * * *', 'Asia/Shanghai')
            """,
            (json_dumps({"schedule": {"kind": "recurring"}}),),
        )
        conn.execute(
            """
            INSERT INTO sent_history(subscription_id, cluster_key, run_id, docx_path)
            VALUES ('sub-1', 'cluster-1', ?, ?)
            """,
            (run_id, str(docx_path)),
        )


if __name__ == "__main__":
    unittest.main()
