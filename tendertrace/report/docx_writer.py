from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tendertrace.adapters.ccgp import Notice
from tendertrace.opportunity import (
    build_market_context,
    competition_context_for_notice,
    intelligence_for_notice,
    market_benchmark_for_notice,
)
from tendertrace.report.naming import safe_report_filename


def _set_run_font(
    run, *, size: float = 11, bold: bool | None = None, color: str | None = None
) -> None:
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _style_doc(doc: Document) -> None:
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
    ):
        heading = doc.styles[name]
        heading.font.name = "Calibri"
        heading._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        heading.font.size = Pt(size)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor.from_string(color)
        heading.paragraph_format.space_before = Pt(before)
        heading.paragraph_format.space_after = Pt(after)


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != 9360:
        raise ValueError("table widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    for tag, attributes in (
        ("w:tblW", {"w:w": "9360", "w:type": "dxa"}),
        ("w:tblInd", {"w:w": "120", "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = table_pr.first_child_found_in(tag)
        if node is None:
            node = OxmlElement(tag)
            table_pr.append(node)
        for key, value in attributes.items():
            node.set(qn(key), value)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _set_run_font(run, size=9, bold=bold)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_border(cell, *, color: str = "D8E0EA", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def _add_label_value(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}：")
    _set_run_font(label_run, bold=True)
    value_run = paragraph.add_run(value or "无")
    _set_run_font(value_run)


def _intelligence(notice: Notice) -> dict[str, Any]:
    return intelligence_for_notice(notice)


def _score_value(intelligence: dict[str, Any], key: str) -> int:
    scores = intelligence.get("scores")
    if not isinstance(scores, dict):
        return 0
    try:
        return int(scores.get(key) or 0)
    except (TypeError, ValueError):
        return 0


def _add_opportunity_summary(doc: Document, notices: list[Notice]) -> None:
    rows = [(notice, _intelligence(notice)) for notice in notices]
    high_priority = sum(1 for _notice, item in rows if item.get("level") == "A")
    average = round(
        sum(int(item.get("score") or 0) for _notice, item in rows) / len(rows),
        1,
    )
    kpis = doc.add_table(rows=1, cols=4)
    kpis.autofit = True
    for cell, (label, value) in zip(
        kpis.rows[0].cells,
        (
            ("公告数量", str(len(rows))),
            ("A 级机会", str(high_priority)),
            ("平均评分", str(average)),
            ("多源佐证", str(sum(1 for notice, _item in rows if int(notice.fields.get("duplicate_count") or 1) > 1))),
        ),
        strict=True,
    ):
        cell.text = ""
        _shade(cell, "F3F7FB")
        _set_cell_border(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = paragraph.add_run(f"{value}\n")
        _set_run_font(value_run, size=16, bold=True, color="174EA6")
        label_run = paragraph.add_run(label)
        _set_run_font(label_run, size=9, color="5B677A")
    _set_table_geometry(kpis, [2340, 2340, 2340, 2340])

    doc.add_heading("机会优先级", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["等级", "综合", "时效", "完整", "可信", "阶段", "项目"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, header, bold=True)
        _shade(cell, "E8EEF5")
    for notice, item in sorted(rows, key=lambda value: int(value[1].get("score") or 0), reverse=True):
        row = table.add_row().cells
        values = [
            f"{item.get('level', 'D')} · {item.get('level_label', '待研判')}",
            str(item.get("score") or 0),
            str(_score_value(item, "freshness")),
            str(_score_value(item, "completeness")),
            str(_score_value(item, "credibility")),
            str(item.get("stage") or "线索识别"),
            notice.title,
        ]
        for cell, value in zip(row, values, strict=True):
            _set_cell_text(cell, value)
    _set_table_geometry(table, [700, 750, 750, 800, 800, 1000, 4560])


def _add_opportunity_detail(
    doc: Document,
    notice: Notice,
    market: dict[str, object],
) -> None:
    intelligence = _intelligence(notice)
    _add_label_value(
        doc,
        "机会研判",
        f"{intelligence.get('level', 'D')} 级 / {intelligence.get('level_label', '待研判')} / "
        f"综合 {intelligence.get('score', 0)} 分 / {intelligence.get('stage', '线索识别')}",
    )
    _add_label_value(
        doc,
        "质量分项",
        f"时效 {_score_value(intelligence, 'freshness')} · "
        f"完整 {_score_value(intelligence, 'completeness')} · "
        f"可信 {_score_value(intelligence, 'credibility')} · "
        f"可行动 {_score_value(intelligence, 'readiness')}",
    )
    _add_label_value(doc, "项目目标", str(intelligence.get("project_target") or "待确认"))
    _add_label_value(doc, "建议策略", str(intelligence.get("strategy") or "待确认"))
    benchmark = market_benchmark_for_notice(notice, market)
    _add_label_value(doc, "市场价格位置", str(benchmark.get("message") or "样本不足"))
    competition = competition_context_for_notice(notice, market)
    _add_label_value(doc, "竞争情报", str(competition.get("message") or "样本不足"))
    evidence_excerpt = str(competition.get("evidence_excerpt") or "").strip()
    if evidence_excerpt:
        _add_label_value(doc, "竞争证据", evidence_excerpt)
    requirement_review = intelligence.get("requirement_review")
    if isinstance(requirement_review, dict):
        _add_label_value(
            doc,
            "需求覆盖",
            f"当前采集文本覆盖 {requirement_review.get('covered_count', 0)}/"
            f"{requirement_review.get('total_count', 0)} 项（"
            f"{requirement_review.get('coverage_score', 0)} 分）",
        )
        missing = requirement_review.get("missing")
        if isinstance(missing, list) and missing:
            _add_label_value(doc, "需求待核对", "、".join(str(item) for item in missing))
        recommendations = requirement_review.get("recommendations")
        if isinstance(recommendations, list) and recommendations:
            heading = doc.add_paragraph("需求优化建议：")
            heading.paragraph_format.keep_with_next = True
            for index, value in enumerate(recommendations):
                paragraph = doc.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = index < len(recommendations) - 1
                _set_run_font(paragraph.add_run(str(value)))
        _add_label_value(doc, "需求研判边界", str(requirement_review.get("basis") or ""))
    risks = intelligence.get("risks")
    if isinstance(risks, list) and risks:
        _add_label_value(doc, "风险提示", "；".join(str(item) for item in risks))
    actions = intelligence.get("recommended_actions")
    if isinstance(actions, list) and actions:
        doc.add_paragraph("角色行动：")
        for item in actions:
            if not isinstance(item, dict):
                continue
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(f"{item.get('role', '负责人')}：{item.get('action', '')}")
            _set_run_font(run)


def _add_market_context(doc: Document, market: dict[str, object]) -> None:
    doc.add_heading("市场研判", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    signals = market.get("signals")
    rows = list(signals) if isinstance(signals, list) else []
    if not rows:
        rows = ["当前结果不足以形成市场基准"]
    for index, value in enumerate(rows, start=1):
        cells = table.add_row().cells
        _set_cell_text(cells[0], f"研判 {index}", bold=True)
        _shade(cells[0], "F3F7FB")
        _set_cell_text(cells[1], str(value))
    _set_table_geometry(table, [1400, 7960])


def _notice_from_dict(value: dict[str, Any]) -> Notice:
    from tendertrace.adapters.ccgp import Attachment

    return Notice(
        id=value["id"],
        source_site=value["source_site"],
        title=value["title"],
        publish_time=value["publish_time"],
        region=value["region"],
        purchaser=value.get("purchaser", ""),
        source_url=value["source_url"],
        content_text=value.get("content_text", ""),
        core_content=value.get("core_content", ""),
        attachments=[Attachment(**item) for item in value.get("attachments", [])],
        fields=value.get("fields", {}),
    )


def _source_display(notice: Notice) -> str:
    value = notice.fields.get("source_sites")
    if isinstance(value, list) and value:
        return "、".join(str(item) for item in value)
    return notice.source_site


def _evidence(notice: Notice) -> dict[str, Any]:
    value = notice.fields.get("evidence")
    return value if isinstance(value, dict) else {}


def _evidence_status_text(notice: Notice) -> str:
    evidence = _evidence(notice)
    status = str(evidence.get("status") or notice.fields.get("evidence_status") or "未校验")
    score = evidence.get("quality_score", notice.fields.get("evidence_score"))
    if score is None:
        return status
    return f"{status}（score: {score}）"


def _attachment_parse_summary(notice: Notice) -> str:
    evidence = _evidence(notice)
    attachments = evidence.get("attachments")
    if not isinstance(attachments, list) or not attachments:
        return "无附件快照"
    extracted = sum(
        1
        for item in attachments
        if isinstance(item, dict) and str(item.get("status") or "") == "extracted"
    )
    downloaded = sum(
        1
        for item in attachments
        if isinstance(item, dict) and str(item.get("status") or "") in {"downloaded", "extracted"}
    )
    failed = sum(
        1
        for item in attachments
        if isinstance(item, dict) and str(item.get("status") or "") == "failed"
    )
    skipped = sum(
        1
        for item in attachments
        if isinstance(item, dict) and str(item.get("status") or "") == "skipped"
    )
    return f"已下载 {downloaded}/{len(attachments)}，已抽取正文 {extracted}，失败 {failed}，跳过 {skipped}"


def _attachment_excerpts(notice: Notice) -> list[str]:
    evidence = _evidence(notice)
    attachments = evidence.get("attachments")
    if not isinstance(attachments, list):
        return []
    excerpts: list[str] = []
    for item in attachments:
        if not isinstance(item, dict):
            continue
        excerpt = str(item.get("text_excerpt") or "").strip()
        if not excerpt:
            continue
        name = str(item.get("name") or item.get("url") or "附件")
        excerpts.append(f"{name}：{excerpt[:360]}")
    return excerpts


def _structured_fields(notice: Notice) -> dict[str, Any]:
    value = notice.fields.get("structured_fields")
    return value if isinstance(value, dict) else {}


def _structured_field_evidence(notice: Notice) -> dict[str, Any]:
    value = notice.fields.get("structured_field_evidence")
    return value if isinstance(value, dict) else {}


def _add_structured_fields(doc: Document, notice: Notice) -> None:
    structured = _structured_fields(notice)
    for label, key in (
        ("项目编号", "project_no"),
        ("预算金额", "budget"),
        ("投标截止", "bid_deadline"),
        ("开标时间", "opening_time"),
    ):
        value = str(structured.get(key) or "").strip()
        if value:
            _add_label_value(doc, label, value)
    _add_structured_evidence(doc, notice)


def _add_structured_evidence(doc: Document, notice: Notice) -> None:
    evidence = _structured_field_evidence(notice)
    if not evidence:
        return
    lines: list[str] = []
    for key in ("project_no", "budget", "bid_deadline", "opening_time"):
        item = evidence.get(key)
        if not isinstance(item, dict):
            continue
        text = str(item.get("evidence_text") or "").strip()
        if text:
            lines.append(f"{key}: {text[:180]}")
    if lines:
        _add_label_value(doc, "字段证据", " | ".join(lines))


def _add_run_stats(doc: Document, run_stats: dict[str, Any] | None) -> None:
    if not run_stats:
        return
    collected = _stat_int(run_stats, "collected")
    deduped = _stat_int(run_stats, "deduped")
    local_retrieved = _stat_int(run_stats, "local_retrieved")
    source_collected = _stat_int(run_stats, "source_collected")
    pre_skipped = _stat_int(run_stats, "pre_skipped_sent")
    doc.add_paragraph(
        f"运行漏斗：候选 {collected} 条，本地复用 {local_retrieved} 条，"
        f"外部来源新增 {source_collected} 条，清洗去重后 {deduped} 条。"
    )
    if pre_skipped:
        doc.add_paragraph(
            f"订阅增量：已跳过历史推送内容 {pre_skipped} 条，本文件仅保留本轮新增内容。"
        )
    source_stats = _source_stats(run_stats)
    if source_stats:
        attempted = len(source_stats)
        hit_sources = sum(1 for item in source_stats if _stat_int(item, "count") > 0)
        failed = sum(1 for item in source_stats if str(item.get("status") or "") == "failed")
        doc.add_paragraph(
            f"多源覆盖：本轮尝试 {attempted} 个来源，{hit_sources} 个来源命中，{failed} 个来源异常。"
        )
        _add_source_coverage_table(doc, source_stats)


def _add_source_coverage_table(doc: Document, source_stats: list[dict[str, Any]]) -> None:
    doc.add_heading("来源覆盖与抓取健康", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["来源", "状态", "命中", "请求", "成功", "阻断", "重试", "说明"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, header, bold=True)
        _shade(cell, "E8EEF5")
    for item in source_stats:
        fetch_stats = item.get("fetch_stats")
        fetch = fetch_stats if isinstance(fetch_stats, dict) else {}
        note_parts = []
        if item.get("relaxed_city"):
            note_parts.append("城市无结果，已放宽到省级")
        if item.get("error"):
            note_parts.append(str(item["error"]))
        if fetch.get("last_error"):
            note_parts.append(str(fetch["last_error"]))
        row = table.add_row().cells
        values = [
            str(item.get("source") or "-"),
            str(item.get("status") or "-"),
            str(item.get("count") or 0),
            str(fetch.get("requests") or 0),
            str(fetch.get("successes") or 0),
            str(fetch.get("blocked") or 0),
            str(fetch.get("retries") or 0),
            "；".join(note_parts) or "正常",
        ]
        for cell, value in zip(row, values, strict=True):
            _set_cell_text(cell, value)
    _set_table_geometry(table, [1100, 900, 650, 700, 700, 700, 700, 3910])


def _source_stats(run_stats: dict[str, Any]) -> list[dict[str, Any]]:
    items = run_stats.get("source_stats")
    if not isinstance(items, list):
        return []
    return [item for item in items if isinstance(item, dict)]


def _stat_int(stats: dict[str, Any], key: str) -> int:
    try:
        return int(stats.get(key) or 0)
    except (TypeError, ValueError):
        return 0


def _region_scope(bidql: dict[str, Any]) -> dict[str, Any]:
    meta = bidql.get("meta")
    if not isinstance(meta, dict):
        return {}
    value = meta.get("region_scope")
    return value if isinstance(value, dict) else {}


def _region_display(bidql: dict[str, Any]) -> str:
    region = bidql.get("region")
    if not isinstance(region, dict):
        return "未限定"
    province = str(region.get("province") or "").strip()
    city = str(region.get("city") or "").strip()
    scope = _region_scope(bidql)
    if scope.get("status") == "relaxed_city" and city and province:
        return f"{city}（城市级无结果，已扩大至{province}省内）"
    if city and province:
        return f"{city}（{province}）"
    return province or "未限定"


def _region_scope_message(bidql: dict[str, Any]) -> str:
    scope = _region_scope(bidql)
    if scope.get("status") != "relaxed_city":
        return ""
    return str(scope.get("message") or "").strip()


def write_report(
    *,
    query: str,
    bidql: dict[str, Any],
    notices: list[Notice] | list[dict[str, Any]],
    output_dir: Path,
    generated_at: datetime,
    run_mode: str = "full",
    run_stats: dict[str, Any] | None = None,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    normalized = [item if isinstance(item, Notice) else _notice_from_dict(item) for item in notices]
    path = output_dir / safe_report_filename(query, generated_at)

    doc = Document()
    _style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("TenderTrace 招投标信息汇总报告")
    _set_run_font(run, size=20, bold=True, color="000000")
    _add_label_value(doc, "用户问题", query)
    _add_label_value(doc, "生成时间", generated_at.strftime("%Y-%m-%d %H:%M"))
    _add_label_value(doc, "运行模式", run_mode)
    _add_label_value(doc, "关键词", "、".join(bidql.get("topic", {}).get("core", [])))
    _add_label_value(doc, "地区", _region_display(bidql))
    window = bidql.get("time", {}).get("resolved_window")
    if window:
        _add_label_value(doc, "时间窗口", f"{window['from']} 至 {window['to']}")
    region_note = _region_scope_message(bidql)
    if region_note:
        _add_label_value(doc, "地域范围说明", region_note)

    doc.add_heading("执行摘要", level=1)
    source_sites = sorted(
        {
            site
            for notice in normalized
            for site in (
                notice.fields.get("source_sites")
                if isinstance(notice.fields.get("source_sites"), list)
                else [notice.source_site]
            )
        }
    )
    source_text = "、".join(source_sites) if source_sites else "已配置来源"
    doc.add_paragraph(f"本次从 {source_text} 筛选出 {len(normalized)} 条符合条件的招投标信息。")
    if region_note:
        doc.add_paragraph(region_note)
    _add_run_stats(doc, run_stats)
    if not normalized:
        doc.add_paragraph("本轮未发现符合条件的新增招投标信息。")
        doc.save(path)
        return path

    market = build_market_context(normalized, as_of=generated_at)
    _add_opportunity_summary(doc, normalized)
    _add_market_context(doc, market)

    doc.add_page_break()
    doc.add_heading("结果总览", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["序号", "标题", "发布时间", "来源", "地区", "附件数"]
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, header, bold=True)
        _shade(cell, "E8EEF5")
    for index, notice in enumerate(normalized, start=1):
        row = table.add_row().cells
        values = [
            str(index),
            notice.title,
            notice.publish_time,
            _source_display(notice),
            notice.region,
            str(len(notice.attachments)),
        ]
        for cell, value in zip(row, values, strict=True):
            _set_cell_text(cell, value)
    _set_table_geometry(table, [600, 3900, 1200, 1100, 1200, 1360])

    doc.add_heading("逐条详情", level=1)
    for index, notice in enumerate(normalized, start=1):
        doc.add_heading(f"{index}. {notice.title}", level=2)
        _add_label_value(doc, "标题", notice.title)
        _add_label_value(doc, "发布时间", notice.publish_time)
        _add_label_value(doc, "来源链接", notice.source_url)
        _add_label_value(doc, "采购人", notice.purchaser)
        _add_label_value(doc, "核心内容", notice.core_content)
        _add_opportunity_detail(doc, notice, market)
        _add_structured_fields(doc, notice)
        evidence = _evidence(notice)
        if evidence:
            _add_label_value(doc, "事实校验", _evidence_status_text(notice))
            _add_label_value(doc, "证据哈希", str(evidence.get("snapshot_sha256") or "无"))
            _add_label_value(doc, "证据摘录", str(evidence.get("excerpt") or "无"))
            _add_label_value(doc, "附件解析", _attachment_parse_summary(notice))
        if int(notice.fields.get("duplicate_count") or 1) > 1:
            _add_label_value(doc, "关联来源", _source_display(notice))
        if notice.attachments:
            doc.add_paragraph("附件链接：")
            for attachment in notice.attachments:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{attachment.name} - {attachment.url}")
                _set_run_font(run)
            excerpts = _attachment_excerpts(notice)
            if excerpts:
                doc.add_paragraph("附件正文摘录：")
                for excerpt in excerpts:
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(excerpt)
                    _set_run_font(run)
        else:
            _add_label_value(doc, "附件链接", "无")

    doc.add_heading("附录", level=1)
    doc.add_paragraph(
        "去重说明：本报告按来源站点、公告 ID、规范化 URL 和 cluster_key 去重；订阅增量另由 sent_history 控制。"
    )
    doc.add_paragraph(
        "证据说明：来源链接指向对应公告详情页；核心内容由详情页正文抽取，不做生成式改写；"
        "事实校验基于标题、详情正文、核心内容和附件链接的可追溯证据生成。"
    )
    doc.save(path)
    return path
