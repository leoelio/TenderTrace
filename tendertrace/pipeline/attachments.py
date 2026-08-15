from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass
from io import BytesIO
import hashlib
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from tendertrace.adapters.ccgp import Notice
from tendertrace.config import Settings
from tendertrace.pipeline.dedup import canonicalize_url, clean_text
from tendertrace.pipeline.evidence import attachment_type
from tendertrace.public_http import UnsafeUrlError, ensure_public_http_url, fetch_public_bytes


DEFAULT_MAX_ATTACHMENTS_PER_NOTICE = 3
DEFAULT_MAX_ATTACHMENT_BYTES = 8 * 1024 * 1024
DEFAULT_ATTACHMENT_TIMEOUT = 8.0
EXTRACTABLE_TYPES = {"pdf", "docx", "xlsx"}

Downloader = Callable[[str, int], bytes]


@dataclass(frozen=True)
class AttachmentSnapshotResult:
    notices: list[Notice]
    stats: dict[str, object]


def enrich_attachment_snapshots(
    notices: list[Notice],
    *,
    settings: Settings,
    max_per_notice: int = DEFAULT_MAX_ATTACHMENTS_PER_NOTICE,
    max_bytes: int = DEFAULT_MAX_ATTACHMENT_BYTES,
    downloader: Downloader | None = None,
) -> AttachmentSnapshotResult:
    download = downloader or _download_bytes
    enriched: list[Notice] = []
    stats = {
        "attachments_seen": 0,
        "attachments_attempted": 0,
        "attachments_downloaded": 0,
        "attachments_extracted": 0,
        "attachments_skipped": 0,
        "attachments_failed": 0,
    }
    for notice in notices:
        records: list[dict[str, Any]] = []
        for index, attachment in enumerate(notice.attachments):
            stats["attachments_seen"] += 1
            if index >= max_per_notice:
                stats["attachments_skipped"] += 1
                records.append(_record_skipped(attachment.name, attachment.url, "max_per_notice"))
                continue
            stats["attachments_attempted"] += 1
            record = _snapshot_attachment(
                name=attachment.name,
                url=attachment.url,
                settings=settings,
                max_bytes=max_bytes,
                downloader=download,
            )
            if record["status"] in {"downloaded", "extracted"}:
                stats["attachments_downloaded"] += 1
            if record["status"] == "extracted":
                stats["attachments_extracted"] += 1
            if record["status"] == "failed":
                stats["attachments_failed"] += 1
            if record["status"] == "skipped":
                stats["attachments_skipped"] += 1
            records.append(record)
        fields = {**notice.fields, "attachment_snapshots": records}
        enriched.append(
            Notice(
                id=notice.id,
                source_site=notice.source_site,
                title=notice.title,
                publish_time=notice.publish_time,
                region=notice.region,
                purchaser=notice.purchaser,
                source_url=notice.source_url,
                content_text=notice.content_text,
                core_content=notice.core_content,
                attachments=notice.attachments,
                fields=fields,
            )
        )
    return AttachmentSnapshotResult(notices=enriched, stats=stats)


def _snapshot_attachment(
    *,
    name: str,
    url: str,
    settings: Settings,
    max_bytes: int,
    downloader: Downloader,
) -> dict[str, Any]:
    try:
        ensure_public_http_url(url, resolve=False)
    except UnsafeUrlError:
        return _record_skipped(name, url, "unsafe_url")
    kind = attachment_type(url)
    if not kind:
        return _record_skipped(name, url, "unknown_type")
    if kind not in EXTRACTABLE_TYPES:
        return _record_skipped(name, url, "unsupported_text_extraction")
    try:
        data = downloader(url, max_bytes)
    except Exception as exc:
        return {
            **_base_record(name, url, kind),
            "status": "failed",
            "error": f"{type(exc).__name__}: {exc}",
        }
    sha256 = hashlib.sha256(data).hexdigest()
    path = _write_attachment(settings.snapshots_dir, sha256, kind, data)
    text = _extract_text(kind, data)
    status = "extracted" if text else "downloaded"
    return {
        **_base_record(name, url, kind),
        "status": status,
        "path": str(path),
        "sha256": sha256,
        "bytes": len(data),
        "text_excerpt": text[:1200],
        "text_length": len(text),
        "error": "" if kind in EXTRACTABLE_TYPES else "unsupported_text_extraction",
    }


def _download_bytes(url: str, max_bytes: int) -> bytes:
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/126 Safari/537.36"
        )
    }
    return fetch_public_bytes(
        url,
        max_bytes=max_bytes,
        timeout=DEFAULT_ATTACHMENT_TIMEOUT,
        headers=headers,
    ).data


def _write_attachment(root: Path, sha256: str, kind: str, data: bytes) -> Path:
    directory = root / "attachments" / sha256[:2]
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / f"{sha256}.{kind}"
    if not path.exists():
        path.write_bytes(data)
    return path


def _extract_text(kind: str, data: bytes) -> str:
    if kind == "pdf":
        return _extract_pdf_text(data)
    if kind == "docx":
        return _extract_docx_text(data)
    if kind == "xlsx":
        return _extract_xlsx_text(data)
    return ""


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages[:8]:
        parts.append(page.extract_text() or "")
    return clean_text(" ".join(parts))


def _extract_docx_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables[:5]:
        for row in table.rows[:30]:
            parts.extend(cell.text for cell in row.cells[:10])
    return clean_text(" ".join(parts))


def _extract_xlsx_text(data: bytes) -> str:
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets[:3]:
        parts.append(sheet.title)
        for row in sheet.iter_rows(max_row=50, max_col=12, values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                parts.append(" ".join(values))
    workbook.close()
    return clean_text(" ".join(parts))


def _record_skipped(name: str, url: str, reason: str) -> dict[str, Any]:
    return {
        **_base_record(name, url, attachment_type(url)),
        "status": "skipped",
        "error": reason,
    }


def _base_record(name: str, url: str, kind: str) -> dict[str, Any]:
    return {
        "name": name,
        "url": url,
        "canonical_url": canonicalize_url(url),
        "type": kind,
        "status": "",
        "path": "",
        "sha256": "",
        "bytes": 0,
        "text_excerpt": "",
        "text_length": 0,
        "error": "",
    }
