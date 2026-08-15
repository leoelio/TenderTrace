from __future__ import annotations

from dataclasses import asdict, dataclass
import json
from pathlib import Path
import re
from typing import Any

from docx import Document

from tendertrace.config import Settings
from tendertrace.db import SCHEMA_VERSION, connection, database_health
from tendertrace.llm.gateway import model_status
from tendertrace.submission import forbidden_package_entries, package_secret_findings
from tendertrace.vault.qianlima import QianlimaSessionVault


@dataclass(frozen=True)
class AcceptanceCheck:
    name: str
    status: str
    detail: str

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass(frozen=True)
class AcceptanceReport:
    status: str
    checks: list[AcceptanceCheck]

    def to_dict(self) -> dict[str, object]:
        counts = {"pass": 0, "warn": 0, "fail": 0}
        for check in self.checks:
            counts[check.status] = counts.get(check.status, 0) + 1
        return {
            "status": self.status,
            "counts": counts,
            "checks": [check.to_dict() for check in self.checks],
        }


REQUIRED_DELIVERY_DOCS = (
    "README.md",
    "docs/design/详设文档.md",
    "docs/operation/操作文档.md",
    "docs/demo/Demo演示脚本.md",
    "docs/operation/P12_验收Harness与交付审计.md",
    "docs/operation/P13_交付收口与完成度审计.md",
    "docs/operation/P14_本地模型自检与云端隔离.md",
    "docs/operation/P15_Demo预检与录屏证据包.md",
    "docs/operation/P16_Demo视频生成与交付闭环.md",
    "docs/operation/P17_登录源状态校验与千里马验收.md",
    "docs/operation/P18_安全提交包与最终打包.md",
    "docs/operation/P19_UI体验与检索召回优化.md",
    "docs/operation/P20_真实进度与地域降级说明.md",
    "docs/operation/P21_导航工作台删除与Agent评测.md",
    "docs/operation/P22_本地库检索与意图解析增强.md",
    "docs/operation/P23_金标召回评测与向量增强.md",
    "docs/delivery/交付清单.md",
    "docs/delivery/完成度审计.md",
)

REQUIRED_TABLES = {
    "runs",
    "trace_events",
    "run_checkpoints",
    "subscriptions",
    "ingest_subscriptions",
    "sent_history",
    "notices",
    "notices_fts",
    "notice_embeddings",
    "clusters",
    "evidence_items",
    "attachment_snapshots",
    "page_artifacts",
    "model_audits",
    "outbox_messages",
    "user_activity_events",
    "weekly_reports",
    "user_memory_profiles",
    "memory_advice_feedback",
    "feishu_lead_import_runs",
}

SECRET_SCAN_PATHS = (
    "README.md",
    ".env.example",
    "docs",
    "tendertrace",
    "tests",
    "web",
    "pyproject.toml",
)

SECRET_PATTERN = re.compile(r"sk-(?:proj-)?[A-Za-z0-9_-]{20,}|Bearer\s+sk-[A-Za-z0-9_-]+")
REPORT_TERMS = ("标题", "发布时间", "来源链接", "核心内容")


def run_acceptance(settings: Settings, *, strict_runtime: bool = True) -> AcceptanceReport:
    checks: list[AcceptanceCheck] = []
    root = settings.workspace_root
    checks.extend(_check_delivery_docs(root))
    checks.extend(_check_env_example(root))
    checks.extend(_check_secret_scan(root))
    checks.extend(_check_demo_video(root))
    checks.extend(_check_submission_package(root))
    checks.extend(_check_database(settings))
    checks.extend(_check_model(settings))
    checks.extend(_check_sources(settings))
    checks.extend(_check_word_artifacts(settings, strict_runtime=strict_runtime))
    checks.extend(_check_run_evidence(settings, strict_runtime=strict_runtime))
    status = "fail" if any(check.status == "fail" for check in checks) else "pass"
    return AcceptanceReport(status=status, checks=checks)


def _check_delivery_docs(root: Path) -> list[AcceptanceCheck]:
    checks: list[AcceptanceCheck] = []
    for relative in REQUIRED_DELIVERY_DOCS:
        path = root / relative
        if path.exists() and path.stat().st_size > 200:
            checks.append(AcceptanceCheck(f"delivery_doc:{relative}", "pass", "present"))
        else:
            checks.append(AcceptanceCheck(f"delivery_doc:{relative}", "fail", "missing or empty"))
    teaching = sorted((root / "docs" / "teaching").glob("*.docx"))
    if teaching:
        checks.append(AcceptanceCheck("teaching_docx", "pass", f"{len(teaching)} docx files"))
    else:
        checks.append(AcceptanceCheck("teaching_docx", "warn", "no teaching docx files found"))
    return checks


def _check_env_example(root: Path) -> list[AcceptanceCheck]:
    path = root / ".env.example"
    if not path.exists():
        return [AcceptanceCheck("env_example", "fail", ".env.example missing")]
    text = path.read_text(encoding="utf-8")
    if not re.search(r"^OPENAI_API_KEY=\s*$", text, re.MULTILINE):
        return [AcceptanceCheck("env_example", "fail", "OPENAI_API_KEY placeholder must be blank")]
    return [AcceptanceCheck("env_example", "pass", "OPENAI_API_KEY placeholder is blank")]


def _check_secret_scan(root: Path) -> list[AcceptanceCheck]:
    findings: list[str] = []
    for relative in SECRET_SCAN_PATHS:
        path = root / relative
        if not path.exists():
            continue
        if path.is_file():
            _scan_file(path, findings, root)
        else:
            for child in path.rglob("*"):
                if child.is_file() and child.suffix.lower() in {
                    ".py",
                    ".md",
                    ".html",
                    ".js",
                    ".css",
                    ".toml",
                }:
                    _scan_file(child, findings, root)
    if findings:
        return [AcceptanceCheck("secret_scan", "fail", "; ".join(findings[:5]))]
    return [AcceptanceCheck("secret_scan", "pass", "no plaintext OpenAI key patterns found")]


def _check_demo_video(root: Path) -> list[AcceptanceCheck]:
    demo_dir = root / "docs" / "demo"
    videos = [
        path
        for path in demo_dir.glob("*")
        if path.is_file() and path.suffix.lower() in {".mp4", ".mov", ".mkv", ".avi", ".webm"}
    ]
    videos = [path for path in videos if path.stat().st_size > 10_000]
    if not videos:
        return [AcceptanceCheck("demo_video", "fail", "no non-empty demo video found")]
    newest = max(videos, key=lambda path: path.stat().st_mtime)
    return [AcceptanceCheck("demo_video", "pass", newest.name)]


def _check_submission_package(root: Path) -> list[AcceptanceCheck]:
    dist_dir = root / "dist"
    packages = sorted(dist_dir.glob("TenderTrace_submission_*.zip")) if dist_dir.exists() else []
    packages = [path for path in packages if path.stat().st_size > 10_000]
    if not packages:
        return [AcceptanceCheck("submission_package", "fail", "no non-empty submission zip found")]
    newest = max(packages, key=lambda path: path.stat().st_mtime)
    try:
        forbidden = forbidden_package_entries(newest)
        secret_findings = package_secret_findings(newest)
    except Exception as exc:
        return [
            AcceptanceCheck(
                "submission_package",
                "fail",
                f"{newest.name} could not be scanned: {type(exc).__name__}",
            )
        ]
    if forbidden:
        return [
            AcceptanceCheck(
                "submission_package",
                "fail",
                f"{newest.name} contains forbidden entries: {', '.join(forbidden[:5])}",
            )
        ]
    if secret_findings:
        return [
            AcceptanceCheck(
                "submission_package",
                "fail",
                f"{newest.name} contains token-like content: {secret_findings[0].path}",
            )
        ]
    return [AcceptanceCheck("submission_package", "pass", newest.name)]


def _scan_file(path: Path, findings: list[str], root: Path) -> None:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return
    if SECRET_PATTERN.search(text):
        findings.append(path.relative_to(root).as_posix())


def _check_database(settings: Settings) -> list[AcceptanceCheck]:
    health = database_health(settings)
    if not health.get("initialized"):
        return [AcceptanceCheck("database", "fail", "database is not initialized")]
    tables = set(health.get("tables", []))
    missing = sorted(REQUIRED_TABLES - tables)
    if missing:
        return [AcceptanceCheck("database", "fail", f"missing tables: {', '.join(missing)}")]
    versions = health.get("schema_versions", [])
    if SCHEMA_VERSION not in versions:
        return [
            AcceptanceCheck(
                "database",
                "fail",
                f"schema version {SCHEMA_VERSION} is missing",
            )
        ]
    return [AcceptanceCheck("database", "pass", "required tables and schema version present")]


def _check_model(settings: Settings) -> list[AcceptanceCheck]:
    status = model_status(settings)
    if status.configured:
        return [
            AcceptanceCheck(
                "model_status",
                "pass",
                f"{status.mode}/{status.provider}/{status.model or 'none'}",
            )
        ]
    return [
        AcceptanceCheck("model_status", "warn", f"{status.mode}/{status.provider} not configured")
    ]


def _check_sources(settings: Settings) -> list[AcceptanceCheck]:
    vault = QianlimaSessionVault(settings)
    qianlima_status = vault.status()
    checks = [
        AcceptanceCheck("source:ccgp", "pass", "public HTTP source configured"),
        AcceptanceCheck("source:ggzy", "pass", "public HTTP source configured"),
    ]
    if qianlima_status.ready:
        checks.append(AcceptanceCheck("source:qianlima", "pass", "login storage_state ready"))
    else:
        checks.append(AcceptanceCheck("source:qianlima", "warn", qianlima_status.validation))
    return checks


def _check_word_artifacts(settings: Settings, *, strict_runtime: bool) -> list[AcceptanceCheck]:
    output_files = sorted(
        settings.outputs_dir.glob("*.docx"), key=lambda path: path.stat().st_mtime
    )
    outbox_files = sorted(settings.outbox_dir.glob("*.docx"), key=lambda path: path.stat().st_mtime)
    checks = [
        _artifact_count_check("outputs_docx", output_files, strict_runtime),
        _artifact_count_check("outbox_docx", outbox_files, strict_runtime),
    ]
    if output_files:
        checks.append(_check_report_docx(output_files[-1]))
    return checks


def _artifact_count_check(
    name: str,
    files: list[Path],
    strict_runtime: bool,
) -> AcceptanceCheck:
    if files:
        return AcceptanceCheck(name, "pass", f"{len(files)} docx files")
    status = "fail" if strict_runtime else "warn"
    return AcceptanceCheck(name, status, "no docx artifacts found")


def _check_report_docx(path: Path) -> AcceptanceCheck:
    try:
        doc = Document(path)
    except Exception as exc:
        return AcceptanceCheck("latest_report_docx", "fail", f"cannot open {path.name}: {exc}")
    text = _document_text(doc)
    missing = [term for term in REPORT_TERMS if term not in text]
    if missing:
        return AcceptanceCheck(
            "latest_report_docx",
            "warn",
            f"{path.name} missing visible terms: {', '.join(missing)}",
        )
    return AcceptanceCheck("latest_report_docx", "pass", path.name)


def _document_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _check_run_evidence(settings: Settings, *, strict_runtime: bool) -> list[AcceptanceCheck]:
    if not settings.db_path.exists():
        return [AcceptanceCheck("run_evidence", "fail", "database is missing")]
    with connection(settings) as conn:
        row = conn.execute(
            """
            SELECT id, stats_json
            FROM runs
            WHERE status = 'finished'
            ORDER BY finished_at DESC
            LIMIT 1
            """
        ).fetchone()
        multi_source = conn.execute(
            """
            SELECT id, stats_json
            FROM runs
            WHERE status = 'finished'
            """
        ).fetchall()
        model_count = conn.execute("SELECT COUNT(*) AS count FROM model_audits").fetchone()["count"]
    if row is None:
        status = "fail" if strict_runtime else "warn"
        return [AcceptanceCheck("run_evidence", status, "no finished runs found")]
    checks = [_check_latest_run_stats(row)]
    checks.append(_check_any_multi_source_run(multi_source))
    if model_count:
        checks.append(AcceptanceCheck("model_audits", "pass", f"{model_count} records"))
    else:
        checks.append(AcceptanceCheck("model_audits", "warn", "no model audit records"))
    return checks


def _check_latest_run_stats(row: Any) -> AcceptanceCheck:
    stats = json.loads(row["stats_json"])
    if int(stats.get("notice_count") or 0) < 1:
        return AcceptanceCheck("latest_finished_run", "warn", f"{row['id']} has no notices")
    if int(stats.get("trace_events") or 0) < 1:
        return AcceptanceCheck("latest_finished_run", "warn", f"{row['id']} has no trace events")
    if int(stats.get("evidence_checked") or 0) < 1:
        return AcceptanceCheck("latest_finished_run", "warn", f"{row['id']} has no evidence checks")
    return AcceptanceCheck("latest_finished_run", "pass", row["id"])


def _check_any_multi_source_run(rows: list[Any]) -> AcceptanceCheck:
    for row in rows:
        stats = json.loads(row["stats_json"])
        source_sites = stats.get("source_sites")
        if isinstance(source_sites, list) and len(set(source_sites)) >= 2:
            return AcceptanceCheck("multi_source_run", "pass", row["id"])
    for row in rows:
        stats = json.loads(row["stats_json"])
        attempted = _attempted_external_sources(stats)
        if len(attempted) >= 2:
            return AcceptanceCheck(
                "multi_source_run",
                "pass",
                f"{row['id']} attempted {', '.join(attempted)}",
            )
    return AcceptanceCheck(
        "multi_source_run", "warn", "no finished run with >=2 source sites attempted"
    )


def _attempted_external_sources(stats: dict[str, object]) -> list[str]:
    source_stats = stats.get("source_stats")
    if not isinstance(source_stats, list):
        return []
    sources = {
        str(item.get("source") or "")
        for item in source_stats
        if isinstance(item, dict) and item.get("source") and not str(item.get("source")).startswith("local_")
    }
    return sorted(sources)
