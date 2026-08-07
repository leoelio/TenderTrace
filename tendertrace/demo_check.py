from __future__ import annotations

from dataclasses import asdict, dataclass
from datetime import datetime
import json
from pathlib import Path
from typing import Any

from docx import Document

from tendertrace.config import Settings
from tendertrace.db import connection
from tendertrace.llm.doctor import model_doctor
from tendertrace.sanitize import sanitize_for_output
from tendertrace.submission import forbidden_package_entries, package_secret_findings
from tendertrace.vault.qianlima import QianlimaSessionVault


@dataclass(frozen=True)
class DemoCheck:
    name: str
    status: str
    detail: str

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass(frozen=True)
class DemoEvidenceReport:
    status: str
    generated_at: str
    checks: list[DemoCheck]
    evidence: dict[str, Any]

    def to_dict(self) -> dict[str, object]:
        counts = {"pass": 0, "warn": 0, "fail": 0}
        for check in self.checks:
            counts[check.status] = counts.get(check.status, 0) + 1
        return {
            "status": self.status,
            "generated_at": self.generated_at,
            "counts": counts,
            "checks": [check.to_dict() for check in self.checks],
            "evidence": self.evidence,
        }


REPORT_TERMS = ("标题", "发布时间", "来源链接", "核心内容")
REQUIRED_TRACE_TOOLS = {
    "intent.rule_parser",
    "adapter.multi.collect",
    "pipeline.clean_dedup",
    "pipeline.evidence_validate",
    "report.docx_writer",
}
OPTIONAL_TRACE_TOOLS = {"llm.intent_enhancer", "pipeline.attachment_extract"}
VIDEO_EXTENSIONS = (".mp4", ".mov", ".mkv", ".avi", ".webm")


def run_demo_check(settings: Settings) -> DemoEvidenceReport:
    evidence = _collect_evidence(settings)
    checks = [
        _check_model(settings, evidence),
        _check_sources(settings, evidence),
        _check_finished_runs(evidence),
        _check_word_outbox(evidence),
        _check_trace_flow(evidence),
        _check_subscription_incremental(evidence),
        _check_video_file(settings, evidence),
        _check_submission_package(settings, evidence),
        _check_ci_config(settings, evidence),
        _check_api_token(settings, evidence),
    ]
    status = "fail" if any(check.status == "fail" for check in checks) else "pass"
    return DemoEvidenceReport(
        status=status,
        generated_at=datetime.now().isoformat(timespec="seconds"),
        checks=checks,
        evidence=evidence,
    )


def write_demo_evidence(report: DemoEvidenceReport, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = sanitize_for_output(report.to_dict())
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _collect_evidence(settings: Settings) -> dict[str, Any]:
    evidence: dict[str, Any] = {
        "outputs_docx": [path.name for path in sorted(settings.outputs_dir.glob("*.docx"))],
        "outbox_docx": [path.name for path in sorted(settings.outbox_dir.glob("*.docx"))],
        "video_files": _video_files(settings.workspace_root),
    }
    if not settings.db_path.exists():
        evidence.update(
            {
                "finished_run_count": 0,
                "distinct_finished_queries": 0,
                "active_subscription_count": 0,
                "sent_history_count": 0,
                "latest_finished_run": None,
                "latest_trace_tools": [],
                "outbox_message_count": 0,
            }
        )
        return evidence
    with connection(settings) as conn:
        run_counts = conn.execute(
            """
            SELECT COUNT(*) AS finished_count,
                   COUNT(DISTINCT original_query) AS distinct_queries
            FROM runs
            WHERE status = 'finished'
            """
        ).fetchone()
        latest_run = conn.execute(
            """
            SELECT id, subscription_id, original_query, status, output_docx_path, stats_json
            FROM runs
            WHERE status = 'finished'
            ORDER BY finished_at DESC
            LIMIT 1
            """
        ).fetchone()
        active_subscriptions = conn.execute(
            "SELECT COUNT(*) AS count FROM subscriptions WHERE status = 'active'"
        ).fetchone()["count"]
        sent_history_count = conn.execute("SELECT COUNT(*) AS count FROM sent_history").fetchone()["count"]
        outbox_message_count = conn.execute("SELECT COUNT(*) AS count FROM outbox_messages").fetchone()["count"]
        trace_tools = []
        if latest_run is not None:
            trace_tools = _trace_tools(conn, latest_run["id"])
    evidence.update(
        {
            "finished_run_count": int(run_counts["finished_count"]),
            "distinct_finished_queries": int(run_counts["distinct_queries"]),
            "active_subscription_count": int(active_subscriptions),
            "sent_history_count": int(sent_history_count),
            "outbox_message_count": int(outbox_message_count),
            "latest_finished_run": _latest_run_dict(latest_run),
            "latest_trace_tools": trace_tools,
        }
    )
    return evidence


def _trace_tools(conn, run_id: str) -> list[str]:
    rows = conn.execute(
        """
        SELECT payload_json
        FROM trace_events
        WHERE run_id = ? AND event_type = 'tool_called'
        ORDER BY seq
        """,
        (run_id,),
    ).fetchall()
    tools: list[str] = []
    for row in rows:
        payload = json.loads(row["payload_json"])
        tool = payload.get("tool")
        if isinstance(tool, str):
            tools.append(tool)
    return tools


def _latest_run_dict(row) -> dict[str, Any] | None:
    if row is None:
        return None
    stats = sanitize_for_output(json.loads(row["stats_json"] or "{}"))
    return {
        "id": row["id"],
        "subscription_id": row["subscription_id"],
        "original_query": row["original_query"],
        "status": row["status"],
        "output_docx_path": row["output_docx_path"],
        "stats": stats,
    }



def _video_files(root: Path) -> list[str]:
    demo_dir = root / "docs" / "demo"
    if not demo_dir.exists():
        return []
    return [
        str(path.relative_to(root))
        for path in sorted(demo_dir.rglob("*"))
        if path.is_file() and path.suffix.lower() in VIDEO_EXTENSIONS
    ]


def _check_model(settings: Settings, evidence: dict[str, Any]) -> DemoCheck:
    report = model_doctor(settings)
    evidence["model_doctor"] = report.to_dict()
    return DemoCheck(
        "model_doctor",
        report.status,
        f"mode={report.checks[0].detail if report.checks else 'unknown'}",
    )


def _check_sources(settings: Settings, evidence: dict[str, Any]) -> DemoCheck:
    vault = QianlimaSessionVault(settings)
    qianlima_status = vault.status()
    evidence["sources"] = [
        {"site": "ccgp", "status": "configured"},
        {"site": "ggzy", "status": "configured"},
        qianlima_status.to_dict() | {"site": "qianlima"},
    ]
    if qianlima_status.ready:
        return DemoCheck("sources", "pass", "ccgp, ggzy and qianlima login state are available")
    return DemoCheck("sources", "warn", f"ccgp and ggzy are configured; qianlima {qianlima_status.validation}")


def _check_finished_runs(evidence: dict[str, Any]) -> DemoCheck:
    run_count = int(evidence["finished_run_count"])
    query_count = int(evidence["distinct_finished_queries"])
    if run_count < 1:
        return DemoCheck("finished_runs", "fail", "no finished runs found")
    if query_count < 2:
        return DemoCheck("finished_runs", "fail", f"only {query_count} distinct finished query")
    return DemoCheck("finished_runs", "pass", f"{run_count} finished runs, {query_count} distinct queries")


def _check_word_outbox(evidence: dict[str, Any]) -> DemoCheck:
    output_count = len(evidence["outputs_docx"])
    outbox_count = len(evidence["outbox_docx"])
    latest = evidence.get("latest_finished_run") or {}
    path_raw = latest.get("output_docx_path")
    if output_count < 1 or outbox_count < 1:
        return DemoCheck("word_outbox", "fail", "outputs/ and outbox/ must both contain docx files")
    if not path_raw:
        return DemoCheck("word_outbox", "fail", "latest finished run has no output_docx_path")
    path = Path(path_raw)
    if not path.exists():
        return DemoCheck("word_outbox", "fail", f"latest output missing: {path.name}")
    missing = _missing_report_terms(path)
    if missing:
        return DemoCheck("word_outbox", "warn", f"{path.name} missing visible terms: {', '.join(missing)}")
    return DemoCheck("word_outbox", "pass", f"{output_count} outputs and {outbox_count} outbox docx files")


def _missing_report_terms(path: Path) -> list[str]:
    try:
        doc = Document(path)
    except Exception:
        return list(REPORT_TERMS)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts)
    return [term for term in REPORT_TERMS if term not in text]


def _check_trace_flow(evidence: dict[str, Any]) -> DemoCheck:
    tools = set(evidence["latest_trace_tools"])
    missing = sorted(REQUIRED_TRACE_TOOLS - tools)
    if missing:
        return DemoCheck("trace_flow", "fail", f"latest run trace missing: {', '.join(missing)}")
    optional_present = sorted(OPTIONAL_TRACE_TOOLS & tools)
    return DemoCheck("trace_flow", "pass", f"required trace tools present; optional={optional_present}")


def _check_subscription_incremental(evidence: dict[str, Any]) -> DemoCheck:
    subscriptions = int(evidence["active_subscription_count"])
    sent_history = int(evidence["sent_history_count"])
    if subscriptions < 1:
        return DemoCheck("subscription_incremental", "fail", "no active subscription found")
    if sent_history < 1:
        return DemoCheck("subscription_incremental", "fail", "sent_history has no incremental records")
    return DemoCheck(
        "subscription_incremental",
        "pass",
        f"{subscriptions} active subscriptions, {sent_history} sent_history records",
    )


def _check_video_file(settings: Settings, evidence: dict[str, Any]) -> DemoCheck:
    files = evidence["video_files"]
    if files:
        return DemoCheck("demo_video_file", "pass", f"{len(files)} video file(s) under docs/demo")
    return DemoCheck(
        "demo_video_file",
        "warn",
        "no demo video file found under docs/demo; record it after rehearsal",
    )


def _check_submission_package(settings: Settings, evidence: dict[str, Any]) -> DemoCheck:
    package = _latest_submission_package(settings.workspace_root)
    if package is None:
        evidence["submission_package"] = {"status": "missing"}
        return DemoCheck(
            "submission_package",
            "warn",
            "no dist/TenderTrace_submission_*.zip found; run package-submission before final delivery",
        )
    try:
        forbidden = forbidden_package_entries(package)
        secret_findings = package_secret_findings(package)
    except Exception as exc:
        evidence["submission_package"] = {
            "status": "invalid",
            "path": str(package),
            "error": f"{type(exc).__name__}: {exc}",
        }
        return DemoCheck("submission_package", "fail", f"package scan failed: {package.name}")
    evidence["submission_package"] = {
        "status": "scanned",
        "path": str(package),
        "forbidden_entry_count": len(forbidden),
        "secret_hit_count": len(secret_findings),
        "secret_findings": [finding.to_dict() for finding in secret_findings],
    }
    if forbidden:
        return DemoCheck(
            "submission_package",
            "fail",
            f"{package.name} contains forbidden entries: {', '.join(forbidden[:3])}",
        )
    if secret_findings:
        return DemoCheck(
            "submission_package",
            "fail",
            f"{package.name} contains {len(secret_findings)} secret-like value(s)",
        )
    return DemoCheck("submission_package", "pass", f"{package.name} passed package safety scan")


def _check_ci_config(settings: Settings, evidence: dict[str, Any]) -> DemoCheck:
    path = settings.workspace_root / ".github" / "workflows" / "ci.yml"
    exists = path.is_file()
    evidence["ci_config"] = {"path": str(path), "exists": exists}
    if not exists:
        return DemoCheck("ci_config", "fail", ".github/workflows/ci.yml is missing")
    return DemoCheck("ci_config", "pass", "GitHub Actions CI workflow is present")


def _check_api_token(settings: Settings, evidence: dict[str, Any]) -> DemoCheck:
    evidence["api_security"] = {
        "app_env": settings.app_env,
        "api_token_configured": settings.api_token_present,
    }
    if settings.api_token_present:
        detail = "API token guard is configured"
    else:
        detail = f"API token guard is disabled for {settings.app_env} mode"
    return DemoCheck("api_security", "pass", detail)


def _latest_submission_package(root: Path) -> Path | None:
    packages = sorted(
        (root / "dist").glob("TenderTrace_submission_*.zip"),
        key=lambda path: (path.stat().st_mtime, path.name),
    )
    return packages[-1] if packages else None
